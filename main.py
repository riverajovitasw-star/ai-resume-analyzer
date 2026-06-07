from fastapi import FastAPI, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse

import fitz
from docx import Document
import os

app = FastAPI()

# =========================
# CORS
# =========================
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# =========================
# EXTRACT PDF TEXT
# =========================
def extract_pdf_text(path):
    text = ""

    try:
        pdf = fitz.open(path)

        for page in pdf:
            text += page.get_text()

        pdf.close()

    except Exception as e:
        print("PDF ERROR:", e)

    return text.lower()

# =========================
# EXTRACT DOCX TEXT
# =========================
def extract_docx_text(path):
    text = ""

    try:
        doc = Document(path)

        for para in doc.paragraphs:
            text += para.text + " "

    except Exception as e:
        print("DOCX ERROR:", e)

    return text.lower()

# =========================
# HOME
# =========================
@app.get("/")
def home():
    return {"message": "Backend Running"}

# =========================
# ANALYZE
# =========================
@app.post("/analyze")
async def analyze_resume(
    file: UploadFile = File(...),
    job_description: str = Form("")
):

    filename = file.filename

    with open(filename, "wb") as buffer:
        buffer.write(await file.read())

    # =========================
    # EXTRACT TEXT
    # =========================
    text = ""

    if filename.endswith(".pdf"):
        text = extract_pdf_text(filename)

    elif filename.endswith(".docx"):
        text = extract_docx_text(filename)

    print("\n======= RESUME TEXT =======")
    print(text)
    print("===========================\n")

    # =========================
    # SKILLS DATABASE
    # =========================
    skills_db = [
        "python",
        "java",
        "sql",
        "machine learning",
        "data analysis",
        "communication",
        "teamwork",
        "leadership",
        "react",
        "next.js",
        "fastapi",
        "html",
        "css",
        "javascript",
        "excel",
        "project management"
    ]

    # =========================
    # FIND SKILLS
    # =========================
    found_skills = []

    for skill in skills_db:
        if skill in text:
            found_skills.append(skill)

    # =========================
    # JOB MATCHING
    # =========================
    matched_skills = []
    missing_skills = []

    jd = job_description.lower()

    for skill in skills_db:

        if skill in jd:

            if skill in text:
                matched_skills.append(skill)

            else:
                missing_skills.append(skill)

    # =========================
    # SCORES
    # =========================
    score = min(len(found_skills) * 12, 100)

    ats_score = min(60 + len(found_skills) * 4, 100)

    total = len(matched_skills) + len(missing_skills)

    if total > 0:
        job_match = int((len(matched_skills) / total) * 100)
    else:
        job_match = 0

    # =========================
    # SUGGESTIONS
    # =========================
    suggestions = []

    if len(found_skills) < 5:
        suggestions.append("Add more technical skills")

    if missing_skills:
        suggestions.append("Add missing job-relevant skills")

    if "projects" not in text:
        suggestions.append("Add projects section")

    if "experience" not in text:
        suggestions.append("Add experience section")

    # =========================
    # GENERATE DOCX
    # =========================
    improved = Document()

    improved.add_heading(
        "Improved Resume Suggestions",
        level=1
    )

    improved.add_heading("Skills Found", level=2)

    for skill in found_skills:
        improved.add_paragraph(skill)

    improved.add_heading("Missing Skills", level=2)

    for skill in missing_skills:
        improved.add_paragraph(skill)

    improved.add_heading("Suggestions", level=2)

    for s in suggestions:
        improved.add_paragraph(s)

    improved.save("improved_resume.docx")

    # =========================
    # RESPONSE
    # =========================
    return {
        "score": score,
        "ats_score": ats_score,
        "job_match": job_match,
        "skills": found_skills,
        "matched_skills": matched_skills,
        "missing_skills": missing_skills,
        "suggestions": suggestions
    }

# =========================
# DOWNLOAD
# =========================
@app.get("/download")
async def download_resume():

    return FileResponse(
        "improved_resume.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename="improved_resume.docx"
    )